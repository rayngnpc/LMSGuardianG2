import os
import smtplib
from datetime import datetime, UTC
from typing import List, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email.message import EmailMessage
from dotenv import load_dotenv
import pytz
import requests
import re
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import sys
from collections import defaultdict
import csv
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import traceback
from docx.enum.table import WD_ALIGN_VERTICAL


# Add parent directory to path for content filter
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))


load_dotenv(override=True)


def fetch_all_links_by_module():
    """Fetch latest scanned links grouped by module and sorted by highest risk score"""
    url = "http://127.0.0.1:8000/scrapedcontents/scan"
    try:
        res = requests.get(url)
        res.raise_for_status()
        all_data = res.json()

        modules_data = {}

        for item in all_data:
            module_id = item.get("module_id")
            if not module_id:
                continue
            if module_id not in modules_data:
                modules_data[module_id] = []
            modules_data[module_id].append(item)

        # Sort each module's links by descending risk_score (None treated as -1)
        for module_id, links in modules_data.items():
            modules_data[module_id] = sorted(
                links,
                key=lambda x: (
                    x.get("risk_score") if x.get("risk_score") is not None else -1
                ),
                reverse=True,
            )

        return modules_data

    except Exception as e:
        print(f"Failed to fetch latest scanned links: {e}")
        return {}


def get_module_details():
    """Get module details from the backend"""
    try:
        res = requests.get("http://127.0.0.1:8000/modules/")
        res.raise_for_status()
        return res.json()
    except Exception as e:
        print(f"Failed to fetch module details: {e}")
        return []


def format_scraped_at(raw_ts: str) -> str:
    """Formats an ISO timestamp to Singapore local time."""

    try:
        dt = datetime.fromisoformat(raw_ts)
        sg = pytz.timezone("Asia/Singapore")
        dt = dt.astimezone(sg)
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return raw_ts or ""


def generateModuleReport(
    ucname: str, moduleCode: str, urls: List[dict], baseUrl: str
) -> str:
    """Generates a DOCX report for a module high risks links and saves it."""

    template_path = os.path.join("scraper", "reportgenerator", "reportChau.docx")
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        if "<unit_id>" in para.text:
            para.text = para.text.replace("<unit_id>", moduleCode)
        if "<uc_name>" in para.text:
            para.text = para.text.replace("<uc_name>", ucname)
        if "<course_url>" in para.text:
            para.text = para.text.replace("<course_url>", baseUrl)

    table = doc.tables[0]
    table.autofit = False  # Ensure fixed column widths

    column_widths = [
        Inches(2.0),  # Link URL
        Inches(2.0),  # Link Location
        Inches(0.4),  # Risk Score
        Inches(1.8),  # Risk Category
        Inches(0.8),  # Detected on
        Inches(2.8),  # Action to be taken
    ]

    for link in urls:
        row = table.add_row().cells
        row[0].text = link.get("url_link", "")

        row[1].text = link.get("content_location", "") or ""
        row[2].text = str(link.get("risk_score", ""))
        reason = "No action required."
        risk_category = link.get("risk_category", "")
        if link.get("is_paywall", False):
            reason = "This is a paywalled content. Please remove it unless access is verified and permitted."
            risk_category = "Paywalled Content"
        row[3].text = risk_category
        row[4].text = format_scraped_at(link.get("scraped_at", ""))

        # Determine action

        if "SITE UNREACHABLE" in (link.get("risk_category") or ""):
            reason = "This is an unreachable domain or broken link. Please rectify it immediately to maintain content availability."
        elif "EDUCATION DOMAIN" in (link.get("risk_category") or ""):
            reason = "This is content hosted on an educational institution's domain. Please ensure content is used only if an agreement exists with the institution."
        elif link.get("risk_score", 0) < 0:
            reason = "This link has a high risk score and is likely to be malicious. Please remove it immediately."

        row[5].text = reason

        # Apply widths, wrapping and vertical alignment
        for idx, cell in enumerate(row):
            cell.width = column_widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # Force Word to wrap text in cell
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(
                qn("w:w"), str(int(column_widths[idx].inches * 1000))
            )  # width in twips
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    # Save to report folder
    sg = pytz.timezone("Asia/Singapore")
    safe_code = datetime.now(sg).strftime("%Y-%m-%d")
    filename = f"{safe_code}_{moduleCode}_report.docx"
    output_dir = os.path.join("scraper", "reportgenerator", "report")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    print(f"[SUCCESS] Report saved to: {output_path}")
    return output_path


def generateCSV(ucname: str, moduleCode: str, urls: List[dict], baseUrl: str) -> str:
    """Generates a CSV report for a module high risk links and saves it."""
    sg = pytz.timezone("Asia/Singapore")
    safe_code = datetime.now(sg).strftime("%Y-%m-%d")
    filename = f"{safe_code}_{moduleCode}_report.csv"
    output_dir = os.path.join("scraper", "reportgenerator", "report")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)

    with open(output_path, mode="w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file, delimiter=",")

        # Header content
        writer.writerow(
            ["LMS Information Security Asssessment Tool Report for Unit Coordinators"]
        )
        writer.writerow([])  # Empty row
        writer.writerow(["Unit Coordinator", ucname])
        writer.writerow(["Module code", moduleCode])
        writer.writerow(["Course Page URL", baseUrl])
        writer.writerow([])  # Empty row

        # Header for links table (bolded in Word, plain in CSV)
        writer.writerow(
            [
                "Link URL",
                "Link Location",
                "Risk Score",
                "Risk Category",
                "Detected on",
                "Action to be taken",
            ]
        )

        # Link entries
        for link in urls:
            url_link = link.get("url_link", "")
            content_location = link.get("content_location", "")
            risk_score = link.get("risk_score", "")
            risk_category = link.get("risk_category", "")

            if link.get("is_paywall", False):
                reason = "This is a paywalled content. Please remove it unless access is verified and permitted."
                risk_category = "Paywalled Content"
            elif "SITE UNREACHABLE" in (risk_category or ""):
                reason = "This is an unreachable domain or broken link. Please rectify it immediately to maintain content availability."
            elif "EDUCATION DOMAIN" in (risk_category or ""):
                reason = "This is content hosted on an educational institution's domain. Please ensure content is used only if an agreement exists with the institution."
            elif risk_score != "" and risk_score < 0:
                reason = "This link has a high risk score and is likely to be malicious. Please remove it immediately."
            else:
                reason = "No action required."

            writer.writerow(
                [
                    url_link,
                    content_location,
                    risk_score,
                    risk_category,
                    link.get("scraped_at", ""),
                    reason,
                ]
            )

    print(f"[SUCCESS] CSV report saved to: {output_path}")
    return output_path


def send_email_with_report(
    to_email: str, doc_path: str, csv_path: str, moduleCode: str, ucname: str
):
    """Sends an email with attached DOCX and CSV reports."""

    EMAIL_ADDRESS = os.getenv("EMAIL_USER")
    EMAIL_PASSWORD = os.getenv("EMAIL_PASS")

    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        print("[ERROR] Email credentials not set in environment.")
        return

    # Validate attachment paths
    if not os.path.exists(doc_path):
        print(f"[ERROR] DOCX report not found: {doc_path}")
        return
    if not os.path.exists(csv_path):
        print(f"[ERROR] CSV report not found: {csv_path}")
        return

    msg = EmailMessage()
    msg["From"] = f"LMS Guardian <{EMAIL_ADDRESS}>"
    msg["To"] = to_email
    msg["Subject"] = f"[ALERT] High-risk links detected in {moduleCode}"
    msg["Reply-To"] = "noreply@example.com"

    body = f"""
Dear {ucname},

We wish to inform you that high-risk external links have been detected on the LMS course site for {moduleCode}. As the Unit Coordinator, your attention is required to review and address the issues identified in the attached report.

Please find both a Word document and CSV version of the report enclosed for your reference.

(This is an automatically generated notification. Please do not reply.)

Best regards,  
LMS Guardian Team
"""
    msg.set_content(body)

    # Attach DOCX report
    try:
        with open(doc_path, "rb") as f:
            msg.add_attachment(
                f.read(),
                maintype="application",
                subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=os.path.basename(doc_path),
            )
        print(f"    Attached DOCX: {os.path.basename(doc_path)}")
    except Exception as e:
        print(f"[ERROR] Failed to attach DOCX: {e}")
        traceback.print_exc()

    # Attach CSV report
    try:
        with open(csv_path, "rb") as f:
            msg.add_attachment(
                f.read(),
                maintype="text",
                subtype="csv",
                filename=os.path.basename(csv_path),
            )
        print(f"    Attached CSV: {os.path.basename(csv_path)}")
    except Exception as e:
        print(f"[ERROR] failed to attach CSV: {e}")
        traceback.print_exc()

    # Send email via Gmail SMTP
    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as smtp:
            # smtp.set_debuglevel(1)  # Enable SMTP debug logging
            smtp.ehlo()
            smtp.starttls()
            smtp.ehlo()
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)

        print(f"    Email sent to {to_email}")
    except Exception as e:
        print(f"[ERROR] Failed to send email: {e}")
        traceback.print_exc()


def generate_and_send_module_reports():
    """Generates and emails risk reports for all modules to their respective Unit Coordinators."""

    # Get module details
    modules = get_module_details()
    module_info = {m["module_id"]: m for m in modules}

    # Fetch all unit coordinators once
    response = requests.get("http://10.51.33.25:8000/unitcoordinator/")
    unit_coordinators = response.json()

    # Get all links grouped by module
    modules_data = fetch_all_links_by_module()

    for module_id in sorted(modules_data.keys()):
        links = modules_data[module_id]

        sorted_links = sorted(
            links,
            key=lambda x: (
                x.get("risk_score", 0) >= 0,
                not x.get("is_paywall", False),
                x.get("risk_score", 0),
            ),
        )

        # Filter to only include risky links (risk_score < 0)
        filtered_links = [
            link for link in sorted_links if link.get("risk_score", 0) < 0
        ]

        if not filtered_links:
            print(f"[INFO] No risky links to report for module {module_id}. Skipping.")
            continue

        print(f"\n  Module {module_id} ({len(filtered_links)} risky links):")
        for link in filtered_links:
            print(
                f"      scraped_id={link['scraped_id']} | "
                f"score={link['risk_score']} | "
                f"paywall={link['is_paywall']} | "
                f"url={link['url_link']}"
            )

        # Module metadata
        module = module_info.get(module_id, {})
        uc_id = module.get("uc_id")
        ucname = "MISSING ERROR"
        to_email = None

        # Find matching UC info
        for uc in unit_coordinators:
            if uc["uc_id"] == uc_id:
                ucname = uc.get("full_name", "MISSING NAME")
                to_email = uc.get("email")
                break

        if not to_email:
            print(f"[ERROR] No email found for uc_id={uc_id} (module_id={module_id})")
            continue

        module_code = f"{module.get('unit_code', '')} {module.get('module_name', '').strip()} {module.get('semester', '')}, {module.get('teaching_period', '')}".strip()
        print(f"[Module CODE]{module_code}")
        base_url = f"http://10.51.33.25/moodle/course/view.php?id={module_id+1}"

        try:
            # Generate reports
            report_path = generateModuleReport(
                ucname, module_code, filtered_links, base_url
            )
            print(f"[GENERATED] Word report: {report_path}")

            csv_path = generateCSV(ucname, module_code, filtered_links, base_url)
            print(f"[GENERATED] CSV report: {csv_path}")

            # Send email with both attachments
            send_email_with_report(
                to_email=to_email,
                doc_path=report_path,
                csv_path=csv_path,
                moduleCode=module_code,
                ucname=ucname,
            )

        except Exception as e:
            print(f"[FAILURE] Failed for module {module_id}: {e}")

    print(f"\n[SUCCESS] All reports processed for {len(modules_data)} modules.")


def generateDocumentReportForSecurityOfficer(all_links):
    """Generates a DOCX report of all scanned links for the Security Officer."""

    if not all_links:
        print("[INFO] No links provided for security officer report.")
        return None

    # Clean input
    all_links = [l for l in all_links if isinstance(l, dict)]

    # Create document
    document = Document()

    # Set landscape layout
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = (
        section.right_margin
    ) = Inches(0.5)

    # Title and intro
    document.add_heading("LMS Security Officer Report", 0)
    document.add_paragraph(
        f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    document.add_paragraph(
        "This document summarizes all scanned external links, grouped by module."
    )

    # Group links by module
    grouped = defaultdict(list)
    for link in all_links:
        module_id = link.get("module_id", "UNKNOWN")
        grouped[module_id].append(link)

    # Define column headers and widths
    headers = [
        "Module ID",
        "URL Link",
        "APA7",
        "LMS Location",
        "Risk Score",
        "Risk Category",
        "Local URL",
        "Scraped ID",
        "Detect At",
    ]
    col_widths = [
        Inches(0.8),
        Inches(2),
        Inches(2),
        Inches(2),
        Inches(0.5),
        Inches(2),
        Inches(2),
        Inches(1),
        Inches(1),
    ]

    for module_id, links in grouped.items():
        if not links:
            continue

        module_code = f"MODULE_{module_id}"
        document.add_heading(f"Module: {module_code} (ID: {module_id})", level=1)

        # Create table
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Disable auto-fit
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        # Add header row
        hdr_cells = table.rows[0].cells
        for idx, cell in enumerate(hdr_cells):
            cell.text = headers[idx]
            cell.width = col_widths[idx]
            for p in cell.paragraphs:
                if p.runs:
                    p.runs[0].font.bold = True

        # Add data rows

        for link in links:
            try:
                row = table.add_row().cells
                formattedtime = format_scraped_at(link.get("scraped_at", ""))

                values = [
                    str(module_id),
                    str(link.get("url_link") or "N/A"),
                    str(link.get("apa7") or "N/A"),
                    str(link.get("content_location") or "N/A"),
                    (
                        str(link.get("risk_score"))
                        if link.get("risk_score") is not None
                        else "N/A"
                    ),
                    str(link.get("risk_category") or "N/A"),
                    str(link.get("localurl") or "N/A"),
                    str(link.get("scraped_id") or "N/A"),
                    str(formattedtime),
                ]
                for idx, cell in enumerate(row):
                    cell.text = values[idx]
                    cell.width = col_widths[idx]
            except Exception as e:
                print(f"[FAILURE] Error writing row: {e}")
                continue

    # Save document
    base_dir = os.path.dirname(os.path.abspath(__file__))  # ~/.../reportgenerator
    output_dir = os.path.join(base_dir, "report")
    filename = f"{output_dir}/SecurityReport_AllModules_{datetime.now().strftime('%Y%m%d')}.docx"
    document.save(filename)
    print(f"[SUCCESS] Report saved: {filename}")
    return filename


def generateCSVForSecurityOfficer(all_links):
    """Generates a CSV report of all scanned links for the Security Officer."""

    base_dir = os.path.dirname(os.path.abspath(__file__))  # ~/.../reportgenerator
    output_dir = os.path.join(base_dir, "report")

    os.makedirs(output_dir, exist_ok=True)

    filename = f"{output_dir}/SecurityReport_AllModules_{datetime.now().strftime('%Y%m%d')}.csv"

    with open(filename, mode="w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)

        # Write header
        writer.writerow(
            [
                "Module ID",
                "URL Link",
                "LMS Location",
                "Risk Score",
                "Risk Category",
                "Local URL",
                "Scraped ID",
                "Scraped At",
            ]
        )

        for link in all_links:
            writer.writerow(
                [
                    link.get("module_id", "UNKNOWN"),
                    link.get("url_link", "N/A"),
                    link.get("content_location", "N/A"),
                    link.get("risk_score", "N/A"),
                    link.get("risk_category", "N/A"),
                    link.get("localurl", "N/A"),
                    link.get("scraped_id", "N/A"),
                    link.get("scraped_at", "N/A"),
                ]
            )

    return filename


def generate_and_send_module_reports_to_security_officer():
    """Generates and emails summary reports of all modules to the Security Officer."""

    # Step 1: Get module metadata
    modules = get_module_details()
    module_info = {m["module_id"]: m for m in modules}

    # Step 3: Get all scanned link data grouped by module
    modules_data = fetch_all_links_by_module()

    if not modules_data or not isinstance(modules_data, dict):
        print("[ERROR] No module data returned — cannot generate report.")
        return

    # Step 4: Flatten and sort all links
    all_links = []
    print(f"\nMODULE LINK SUMMARY:")
    for module_id in sorted(modules_data.keys()):
        links = modules_data[module_id]
        if not links:
            continue

        sorted_links = sorted(
            links,
            key=lambda x: (
                (
                    0
                    if x.get("is_paywall", False)
                    else 1 if x.get("risk_score", 0) < 0 else 2
                ),
                x.get("risk_score", 0),
            ),
        )
        all_links.extend(sorted_links)

        scores = [l.get("risk_score", 0) for l in links]
        max_risk = max(scores) if scores else "N/A"
        print(
            f"   Module {module_id}: {len(links)} links, Highest Risk Score: {max_risk}"
        )

    print(f"   Total links: {len(all_links)}")

    # Step 5: Generate reports
    try:

        report_path = generateDocumentReportForSecurityOfficer(all_links)
        print(f"[GENERATED] Word report: {report_path}")

        csv_path = generateCSVForSecurityOfficer(all_links)
        print(f"[GENERATED] CSV report: {csv_path}")

        ucname = "Murdoch Security Officer"
        to_email = "syafiqwork2023@gmail.com"

        # Step 6: Send email with both reports
        send_email_with_report(
            to_email=to_email,
            doc_path=report_path,
            csv_path=csv_path,
            moduleCode="ALL MODULES",
            ucname=ucname,
        )

    except Exception as e:
        print(f"[FAILURE] Failed to generate/send security officer report: {e}")

    print(
        f"\n [SUCCESS] Full security report completed for {len(modules_data)} modules."
    )


# === MAIN TEST ===
if __name__ == "__main__":
    # Generate and send module-specific reports
    print("     Starting module-specific report generation...")
    # for all UCs
    generate_and_send_module_reports()
    # for ISO next
    generate_and_send_module_reports_to_security_officer()
